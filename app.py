import os
from flask import Flask, render_template, send_file
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches

app = Flask(__name__)

# Data sesuai gambar instruksi
DATA = {
    "nama": "Nama Lengkap Anda",
    "profesi": "Backend Developer | Python",
    "tagline": "Membangun sistem yang efisien dan skalabel",
    "ringkasan": "Ringkasan singkat tentang profil profesional saya.",
    "tentang": {
        "deskripsi": "Saya adalah pengembang backend yang fokus pada Python.",
        "minat": "AI, Cloud Computing, Open Source",
        "keahlian_umum": "Arsitektur Microservices, REST API"
    },
    "pendidikan": "S1 Informatika - Universitas ABC (2020 - 2024)",
    "pengalaman": "Magang Backend di Startup X - Mengembangkan API Flask.",
    "keahlian": {
        "hard": "Python, Django, SQL, Flask",
        "soft": "Komunikasi, Teamwork, Leadership"
    },
    "kontak": {
        "email": "email@contoh.com",
        "wa": "62812345678",
        "github": "github.com/username"
    },
    "foto_path": "static/foto.jpg"
}

@app.route('/')
def home():
    return render_template('index.html', data=DATA)

@app.route('/download/pdf')
def download_pdf():
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    if os.path.exists(DATA["foto_path"]):
        p.drawImage(DATA["foto_path"], 450, 650, width=100, height=120)
    p.setFont("Helvetica-Bold", 16)
    p.drawString(50, 750, DATA["nama"])
    p.setFont("Helvetica", 12)
    p.drawString(50, 730, DATA["profesi"])
    p.line(50, 720, 550, 720)
    p.drawString(50, 680, f"Pendidikan: {DATA['pendidikan']}")
    p.drawString(50, 660, f"Keahlian: {DATA['keahlian']['hard']}")
    p.showPage()
    p.save()
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="CV.pdf")

@app.route('/download/word')
def download_word():
    doc = Document()
    if os.path.exists(DATA["foto_path"]):
        doc.add_picture(DATA["foto_path"], width=Inches(1.5))
    doc.add_heading(DATA["nama"], 0)
    doc.add_paragraph(DATA["profesi"], style='Subtitle')
    doc.add_heading('Pendidikan', level=1)
    doc.add_paragraph(DATA["pendidikan"])
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="CV.docx")

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)